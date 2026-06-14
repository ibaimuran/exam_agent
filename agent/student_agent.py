"""学生错题诊断 Agent — 错题归档 + AI 分析 + Word 生成"""
import os
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm

from .base import BaseAgent
from config import PAGE_LAYOUT, TYPOGRAPHY
from prompts import get_student_analysis_system, get_student_analysis_user
from utils.docx_utils import setup_page, set_default_font


class StudentAgent(BaseAgent):
    """学生错题诊断 Agent"""

    def __init__(self, api_key: str = None, base_url: str = None, data_dir: str = None):
        super().__init__(api_key, base_url)
        self.data_dir = data_dir or os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "students"
        )

    def archive(
        self,
        student_name: str,
        wrong_nums: str,
        paper_name: str,
        paper_analysis: str,
        paper_questions: list,
        output_dir: str,
    ) -> dict:
        if not student_name or not wrong_nums:
            return {"error": "请提供姓名和错题号", "step": "validate"}

        wrong_list = self._parse_wrong_nums(wrong_nums)
        wrong_details = self._match_questions(wrong_list, paper_questions)

        system_prompt = get_student_analysis_system()
        user_prompt = get_student_analysis_user(
            student_name, paper_name, wrong_list, paper_analysis, wrong_details
        )
        ai_output = self._call_ai(system_prompt, user_prompt, temp=0.1, max_tok=1500)

        overview, detail_analysis, weakness, measures = self._parse_ai_output(ai_output)

        student = self._load_student(student_name)
        student["records"].append({
            "date": datetime.now().isoformat(),
            "paper": paper_name,
            "wrong_nums": ", ".join(wrong_list),
            "overview": overview,
            "detail_analysis": detail_analysis,
            "weakness": weakness,
            "measures": measures,
            "details": wrong_details,
        })
        self._save_student(student)

        docx_path = self._generate_docx(student_name, student["records"][-1], output_dir)

        return {
            "student_name": student_name,
            "overview": overview,
            "detail_analysis": detail_analysis,
            "weakness": weakness,
            "measures": measures,
            "record_count": len(student["records"]),
            "docx_path": docx_path,
            "step": "done",
        }

    def _parse_wrong_nums(self, wrong_nums: str) -> list:
        result = []
        for item in re.split(r'[,，\s]+', wrong_nums):
            item = item.strip()
            if not item:
                continue
            m = re.match(r'(\d+)((?:\(\d+\))*)', item)
            if m:
                main_q = m.group(1)
                sub_parts = re.findall(r'\((\d+)\)', m.group(2))
                if sub_parts:
                    for sp in sub_parts:
                        result.append(f"{main_q}({sp})")
                else:
                    result.append(main_q)
        return result

    def _match_questions(self, wrong_list: list, paper_questions: list) -> list:
        wrong_details = []
        for w in wrong_list:
            found = False
            for q in paper_questions:
                qid = str(q.get("题号", "")).strip()
                if qid == w:
                    wrong_details.append(q)
                    found = True
                    break
            if not found:
                for q in paper_questions:
                    qid = str(q.get("题号", "")).strip()
                    if w in qid or qid in w:
                        wrong_details.append(q)
                        found = True
                        break
            if not found:
                wrong_details.append({
                    "题号": w, "题型": "未知", "分值": "未知",
                    "知识点": "（请从试卷分析中查找）",
                })
        return wrong_details

    def _parse_ai_output(self, ai_output: str) -> tuple:
        parts = ai_output.split("【")
        overview = next(
            (p.replace("整体表现分析】", "").strip() for p in parts if "整体" in p), ""
        )
        detail = next(
            (p.replace("逐题错因分析】", "").strip() for p in parts if "逐题" in p or "错因" in p), ""
        )
        weakness = next(
            (p.replace("薄弱知识点汇总】", "").strip() for p in parts if "薄弱" in p), ""
        )
        measures = next(
            (p.replace("解决措施】", "").strip() for p in parts if "解决" in p), ""
        )
        return overview, detail, weakness, measures

    def _load_student(self, name: str) -> dict:
        os.makedirs(self.data_dir, exist_ok=True)
        path = os.path.join(self.data_dir, f"{name}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        return {"name": name, "records": []}

    def _save_student(self, data: dict):
        os.makedirs(self.data_dir, exist_ok=True)
        path = os.path.join(self.data_dir, f"{data['name']}.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def _generate_docx(self, student_name: str, record: dict, output_dir: str) -> str:
        doc = Document()
        setup_page(doc)
        set_default_font(doc)

        doc.add_heading(f"{student_name} 试卷分析报告", level=1)

        doc.add_heading("基本信息", level=2)
        info = [
            ("学生姓名", student_name),
            ("试卷", record.get("paper", "")),
            ("错题号", record.get("wrong_nums", "")),
            ("分析日期", datetime.now().strftime("%Y-%m-%d")),
        ]
        for k, v in info:
            p = doc.add_paragraph()
            run = p.add_run(f"{k}：{v}")
            run.font.size = Pt(11)

        doc.add_heading("整体表现分析", level=2)
        doc.add_paragraph(record.get("overview", ""))

        doc.add_heading("逐题错因分析", level=2)
        detail = record.get("detail_analysis", "")
        if detail:
            for line in detail.strip().split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

        doc.add_heading("错题详情", level=2)
        table = doc.add_table(rows=1, cols=4, style="Table Grid")
        hdr = table.rows[0].cells
        for i, h in enumerate(["题号", "题型", "分值", "知识点"]):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for q in record.get("details", []):
            row = table.add_row().cells
            row[0].text = q.get("题号", "")
            row[1].text = q.get("题型", "")
            row[2].text = q.get("分值", "")
            row[3].text = q.get("知识点", "")
        for row in table.rows:
            row.cells[0].width = Cm(2)
            row.cells[1].width = Cm(2)
            row.cells[2].width = Cm(1.5)
            row.cells[3].width = Cm(10)

        doc.add_heading("薄弱知识点汇总", level=2)
        doc.add_paragraph(record.get("weakness", ""))

        doc.add_heading("解决措施", level=2)
        doc.add_paragraph(record.get("measures", ""))

        path = os.path.join(output_dir, f"{student_name}_试卷分析.docx")
        doc.save(path)
        return path

    def load_student(self, name: str) -> dict:
        """公开方法：加载学生数据（供 orchestrator 使用）"""
        return self._load_student(name)
