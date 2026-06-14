"""试卷分析 Agent — PDF/Word 提取 + AI 逐题分析 + Word 生成"""
import os
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base import BaseAgent
from config import PAGE_LAYOUT, PAPER_ANALYSIS_TEMPLATE
from prompts import get_paper_analysis_system, get_paper_analysis_user
from utils.file_utils import extract_text
from utils.docx_utils import setup_page, set_default_font


class PaperAgent(BaseAgent):
    """试卷分析 Agent"""

    def analyze(self, file, paper_name: str, output_dir: str) -> dict:
        result = extract_text(file)
        text = result["text"]
        diag = result.get("diag") or {}

        if text is None:
            return {"error": "仅支持 PDF 和 Word 文件", "step": "extract"}

        if not text.strip():
            pages = diag.get("pages", 0)
            has_img = diag.get("has_images", False)
            if has_img and pages > 0:
                msg = (
                    f"该 PDF 为扫描图片型文件（共 {pages} 页），无文字层可供提取。"
                    "已尝试 Tesseract 和 EasyOCR 识别，请确保已安装 Tesseract OCR 引擎，"
                    "或将 PDF 转换为含文字层的格式后重试。"
                )
            elif pages > 0:
                msg = (
                    f"PDF 共 {pages} 页，但所有提取方式（含 OCR）均未获取到文字。"
                    "文件可能使用了特殊编码或加密，请尝试用 Word 格式上传。"
                )
            else:
                msg = "无法从文件中提取文字。请确认文件未损坏且包含可选中的文字。"
            return {"error": msg, "step": "extract"}

        system_prompt = get_paper_analysis_system()
        user_prompt = get_paper_analysis_user(paper_name, text)
        analysis = self._call_ai(system_prompt, user_prompt, temp=0.1, max_tok=4000)

        questions = self._parse_questions(analysis)
        docx_path = self._generate_docx(paper_name, questions, output_dir)

        return {
            "paper_name": paper_name,
            "analysis": analysis,
            "questions": questions,
            "question_count": len(questions),
            "docx_path": docx_path,
            "raw": text,
            "step": "done",
        }

    def _parse_questions(self, analysis: str) -> list:
        questions = []
        for line in analysis.strip().split("\n"):
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if line.startswith("|"):
                line = line.strip("|")
            parts = [p.strip() for p in line.split("|")]
            if len(parts) >= 4 and parts[0].strip():
                qid = parts[0].strip()
                if qid.isdigit():
                    questions.append({
                        "题号": qid,
                        "题型": parts[1].strip(),
                        "分值": parts[2].strip(),
                        "知识点": parts[3].strip(),
                    })
        return questions

    def _generate_docx(self, paper_name: str, questions: list, output_dir: str) -> str:
        doc = Document()
        setup_page(doc)
        set_default_font(doc)

        title = doc.add_heading(f"{paper_name} 试卷分析", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        hdrs = PAPER_ANALYSIS_TEMPLATE["headers"]
        widths = PAPER_ANALYSIS_TEMPLATE["col_widths"]
        table = doc.add_table(rows=1, cols=len(hdrs), style="Table Grid")
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(hdrs):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True

        for q in questions:
            row = table.add_row().cells
            for ci, k in enumerate(hdrs):
                row[ci].text = q.get(k, "")

        for row in table.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = w

        doc.add_paragraph("")
        doc.add_paragraph(
            f"共 {len(questions)} 题 | 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )

        path = os.path.join(output_dir, f"{paper_name}_试卷分析.docx")
        doc.save(path)
        return path
