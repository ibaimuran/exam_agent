"""共享 Word 文档生成工具"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from config import TYPOGRAPHY, PAGE_LAYOUT


def setup_page(doc: Document):
    """设置 A4 页面"""
    for section in doc.sections:
        section.page_width = PAGE_LAYOUT["width"]
        section.page_height = PAGE_LAYOUT["height"]
        section.left_margin = PAGE_LAYOUT["left_margin"]
        section.right_margin = PAGE_LAYOUT["right_margin"]
        section.top_margin = PAGE_LAYOUT["top_margin"]
        section.bottom_margin = PAGE_LAYOUT["bottom_margin"]


def set_default_font(doc: Document):
    """设置默认字体和行距"""
    style = doc.styles["Normal"]
    style.font.name = TYPOGRAPHY["font_ascii"]
    style.font.size = TYPOGRAPHY["body_size"]
    style.element.rPr.rFonts.set(qn("w:eastAsia"), TYPOGRAPHY["font_cjk"])
    style.paragraph_format.line_spacing = TYPOGRAPHY["line_spacing"]


def add_title(
    doc: Document,
    text: str,
    size=None,
    bold: bool = True,
    center: bool = True,
):
    """添加标题段落"""
    if size is None:
        size = TYPOGRAPHY["title_size"]
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = TYPOGRAPHY["font_ascii"]
    run.element.rPr.rFonts.set(qn("w:eastAsia"), TYPOGRAPHY["font_cjk"])
    run.font.size = size
    run.bold = bold


def add_body(
    doc: Document,
    text: str,
    size=None,
    bold: bool = False,
    indent: bool = False,
):
    """添加正文段落"""
    if size is None:
        size = TYPOGRAPHY["body_size"]
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = TYPOGRAPHY["font_ascii"]
    run.element.rPr.rFonts.set(qn("w:eastAsia"), TYPOGRAPHY["font_cjk"])
    run.font.size = size
    run.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)


def add_docx_table(
    doc: Document,
    headers: list,
    rows: list,
    col_widths: list = None,
):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            if ci < len(table.rows[ri + 1].cells):
                table.rows[ri + 1].cells[ci].text = str(val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    return table
