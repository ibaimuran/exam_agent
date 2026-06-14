"""PDF/Word 文件文字提取"""
import os
import logging
import tempfile
import threading

# 国内 HuggingFace 不可用，使用镜像站下载 OCR 模型
if not os.environ.get("HF_ENDPOINT"):
    os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

logger = logging.getLogger(__name__)

_easyocr_reader = None
_easyocr_lock = threading.Lock()


def init_ocr():
    """预加载 CnOCR 模型（在应用启动时调用，避免首次请求超时）

    EasyOCR 不在此预加载（模型下载慢），改为按需延迟加载。
    """
    try:
        from cnocr import CnOcr
        from PIL import Image
        logger.info("正在加载 CnOCR 模型（仅首次下载，后续秒级加载）...")
        ocr = CnOcr(det_model_name="ch_PP-OCRv3_det", rec_model_name="ch_PP-OCRv3")
        # 用一张小图触发首次模型加载/下载
        img = Image.new("RGB", (100, 40), "white")
        ocr.ocr(img)
        logger.info("CnOCR 模型就绪")
    except ImportError:
        logger.info("cnocr 未安装，OCR 功能受限")
    except Exception as e:
        logger.warning("CnOCR 模型加载失败: %s", e)


def _pdf_diagnostics(file_path: str) -> dict:
    """收集 PDF 文件诊断信息"""
    diag = {"pages": 0, "has_images": False, "size_kb": 0}
    try:
        diag["size_kb"] = round(os.path.getsize(file_path) / 1024, 1)
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        diag["pages"] = len(reader.pages)
        # 粗略判断是否含图片：检查页面是否有 /XObject
        for page in reader.pages:
            if "/XObject" in page.get("/Resources", {}):
                resources = page["/Resources"]
                if "/XObject" in resources:
                    for obj in resources["/XObject"].values():
                        if obj.get("/Subtype") == "/Image":
                            diag["has_images"] = True
                            break
            if diag["has_images"]:
                break
    except Exception:
        pass
    return diag


def _extract_with_pdfplumber(file_path: str) -> str:
    """方案 1: pdfplumber"""
    import pdfplumber
    parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)


def _extract_with_pymupdf(file_path: str) -> str:
    """方案 2: PyMuPDF (fitz) — 对 CJK 和特殊编码支持好"""
    import fitz
    parts = []
    doc = fitz.open(file_path)
    for page in doc:
        t = page.get_text()
        if t and t.strip():
            parts.append(t.strip())
    doc.close()
    return "\n".join(parts)


def _extract_with_pypdf2(file_path: str) -> str:
    """方案 3: PyPDF2"""
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _ocr_with_tesseract(file_path: str, diag: dict) -> str:
    """OCR 方案 A: Tesseract — 速度快，需系统安装 Tesseract 引擎"""
    try:
        import fitz
        from PIL import Image
        import pytesseract
        import io

        doc = fitz.open(file_path)
        parts = []
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            if text and text.strip():
                parts.append(text.strip())
            logger.info("Tesseract OCR: 第 %d/%d 页完成", i + 1, len(doc))
        doc.close()
        if parts:
            logger.info("Tesseract OCR 识别成功，共 %d 页", len(parts))
        return "\n".join(parts)
    except ImportError:
        logger.info("pytesseract 未安装，跳过")
        return ""
    except Exception as e:
        logger.warning("Tesseract OCR 失败: %s", e)
        return ""


def _ocr_with_cnocr(file_path: str, diag: dict) -> str:
    """OCR 方案 B: CnOCR — 专为中文优化，模型从国内源下载，速度快"""
    try:
        import fitz
        from PIL import Image
        from cnocr import CnOcr
        import io

        ocr = CnOcr(det_model_name="ch_PP-OCRv3_det", rec_model_name="ch_PP-OCRv3")
        doc = fitz.open(file_path)
        parts = []
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            results = ocr.ocr(img)
            # CnOCR 返回 [{'text': ..., 'score': ..., 'position': ...}, ...]
            texts = [item["text"] for item in results if item.get("text")]
            page_text = " ".join(texts)
            if page_text.strip():
                parts.append(page_text.strip())
            logger.info("CnOCR: 第 %d/%d 页完成", i + 1, len(doc))
        doc.close()
        if parts:
            logger.info("CnOCR 识别成功，共 %d 页", len(parts))
        return "\n".join(parts)
    except ImportError:
        logger.info("cnocr 未安装，跳过")
        return ""
    except Exception as e:
        logger.warning("CnOCR 失败: %s", e)
        return ""


def _ocr_with_easyocr(file_path: str, diag: dict) -> str:
    """OCR 方案 C: EasyOCR — 纯 Python，无需系统依赖，支持中英文"""
    global _easyocr_reader
    try:
        import fitz
        from PIL import Image
        import io
        import numpy as np

        if _easyocr_reader is None:
            with _easyocr_lock:
                if _easyocr_reader is None:
                    import easyocr
                    _easyocr_reader = easyocr.Reader(["ch_sim", "en"], gpu=False)

        doc = fitz.open(file_path)
        parts = []
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            arr = np.array(img)
            results = _easyocr_reader.readtext(arr, detail=0)
            page_text = " ".join(results)
            if page_text.strip():
                parts.append(page_text.strip())
            logger.info("EasyOCR: 第 %d/%d 页完成", i + 1, len(doc))
        doc.close()
        if parts:
            logger.info("EasyOCR 识别成功，共 %d 页", len(parts))
        return "\n".join(parts)
    except ImportError:
        logger.info("easyocr 未安装，跳过")
        return ""
    except Exception as e:
        logger.warning("EasyOCR 失败: %s", e)
        return ""


def _extract_pdf_text(file_path: str) -> tuple:
    """从 PDF 提取文字，返回 (text, diag)

    依次尝试：pdfplumber → PyMuPDF → PyPDF2 → Tesseract OCR → EasyOCR
    """
    diag = _pdf_diagnostics(file_path)
    logger.info("PDF 诊断: %s 页, %.1f KB, 含图片=%s",
                diag["pages"], diag["size_kb"], diag["has_images"])

    # 阶段 1: 文字提取
    text_methods = [
        ("pdfplumber", _extract_with_pdfplumber),
        ("PyMuPDF",  _extract_with_pymupdf),
        ("PyPDF2",   _extract_with_pypdf2),
    ]

    for name, fn in text_methods:
        try:
            text = fn(file_path)
            if text.strip():
                logger.info("%s 提取成功，共 %d 字符", name, len(text))
                return text, diag
            else:
                logger.info("%s: 未提取到文字", name)
        except ImportError:
            logger.info("%s: 未安装，跳过", name)
        except Exception as e:
            logger.warning("%s 提取失败: %s", name, e)

    # 阶段 2: OCR（图片型 PDF）
    logger.info("文字提取均失败，尝试 OCR ...")
    for name, fn in [
        ("Tesseract", _ocr_with_tesseract),
        ("CnOCR",    _ocr_with_cnocr),
        ("EasyOCR",  _ocr_with_easyocr),
    ]:
        try:
            text = fn(file_path, diag)
            if text.strip():
                logger.info("%s 识别成功，共 %d 字符", name, len(text))
                diag["method"] = name.lower()
                return text, diag
        except Exception:
            pass

    logger.warning("所有提取方式（含 OCR）均未获取到文字")
    return "", diag


def _extract_docx_text(file, tmp_path: str) -> str:
    """从 Word 文件中提取文字（含表格内容）"""
    from docx import Document

    file.save(tmp_path)
    doc = Document(tmp_path)

    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    if not parts:
        for p in doc.paragraphs:
            parts.append(p.text)

    return "\n".join(parts)


def extract_text(file) -> dict:
    """从上传文件中提取文字，支持 PDF 和 Word

    返回:
        {"text": str|None, "diag": dict|None}
        text=None 表示不支持的文件类型
    """
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        text, diag = _extract_pdf_text(tmp_path)
        os.unlink(tmp_path)
        return {"text": text, "diag": diag}

    elif filename.endswith((".docx", ".doc")):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_path = tmp.name
        try:
            text = _extract_docx_text(file, tmp_path)
        finally:
            os.unlink(tmp_path)
        return {"text": text, "diag": None}

    return {"text": None, "diag": None}
